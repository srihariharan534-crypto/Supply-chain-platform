import pandas as pd
from docx import Document
from datetime import datetime

class ReportGenerator:
    def __init__(self, data: dict):
        self.data = data
        
    def generate_executive_report(self, filepath="executive_report.docx"):
        """Generate Executive Report Template."""
        doc = Document()
        doc.add_heading('Executive Supply Chain Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
        
        doc.add_heading('Overall Health', level=1)
        doc.add_paragraph(f"Perfect Order Rate: {self.data.get('perfect_order_rate', 'N/A')}")
        
        doc.add_heading('Critical Alerts', level=1)
        doc.add_paragraph("Review the logistics and inventory anomalies detected by the system.")
        
        doc.save(filepath)
        print(f"Executive report saved to {filepath}")
        
    def generate_inventory_report(self, filepath="inventory_report.docx"):
        """Generate Inventory Report Template."""
        doc = Document()
        doc.add_heading('Inventory Control Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
        doc.save(filepath)

    def generate_logistics_report(self, filepath="logistics_report.docx"):
        """Generate Logistics Report Template."""
        doc = Document()
        doc.add_heading('Logistics Performance Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
        doc.save(filepath)

    def generate_supplier_report(self, filepath="supplier_report.docx"):
        """Generate Supplier Report Template."""
        doc = Document()
        doc.add_heading('Supplier Risk Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
        doc.save(filepath)
        
    def generate_warehouse_report(self, filepath="warehouse_report.docx"):
        """Generate Warehouse Report Template."""
        doc = Document()
        doc.add_heading('Warehouse Utilization Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
        doc.save(filepath)

    def generate_risk_assessment_report(self, filepath="risk_report.docx"):
        """Generate Risk Assessment Report Template."""
        doc = Document()
        doc.add_heading('Comprehensive Risk Assessment', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
        doc.save(filepath)
